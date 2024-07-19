import io
import json
import os

import openpyxl
import pytest
from django.urls import reverse
from docx import Document
from lxml import etree
from rest_framework import status

from document_merge_service.api.data import django_file

from .. import models, serializers


@pytest.mark.parametrize("template__description", ["test description"])
@pytest.mark.parametrize(
    "query_params,size",
    [
        ({"description__icontains": "test"}, 1),
        ({"description__search": "test"}, 1),
        ({"description__icontains": "unknown"}, 0),
        ({"description__search": "unknown"}, 0),
    ],
)
def test_template_list_query_params(db, admin_client, template, size, query_params):
    url = reverse("template-list")

    response = admin_client.get(url, data=query_params)
    assert response.status_code == status.HTTP_200_OK
    assert response.json()["count"] == size


def test_template_detail(db, client, template, snapshot):
    url = reverse("template-detail", args=[template.pk])

    response = client.get(url)
    assert response.status_code == status.HTTP_200_OK
    snapshot.assert_match(response.json())


def test_template_download(db, client, template):
    file = django_file("docx-template-syntax.docx")
    template.template.save(os.path.basename(file.name), file)
    template.save()

    url = reverse("template-detail", args=[template.pk])
    response = client.get(url)

    data = response.json()

    download_url = data["template"]

    template_resp = client.get(download_url)

    file.seek(0)
    assert file.read() == template_resp.getvalue()


def test_template_list_with_file(db, client, template):
    file = django_file("docx-template-syntax.docx")
    template.template.save(os.path.basename(file.name), file)
    template.save()

    url = reverse("template-list")
    response = client.get(url)

    assert response.json()["results"][0]["template"] is not None
    assert response.status_code == status.HTTP_200_OK


def test_template_download_url(db, client, template):
    file = django_file("docx-template-syntax.docx")
    template.template.save(os.path.basename(file.name), file)
    template.save()

    serializer = serializers.TemplateSerializer(template)
    field = serializer.fields["template"]
    assert (
        field.to_representation(template)
        == f"/api/v1/template-download/{template.slug}"
    )


@pytest.mark.parametrize(
    "template_name,engine,status_code,require_authentication,authenticated",
    [
        (
            "xlsx-template.xlsx",
            models.Template.XLSX_TEMPLATE,
            status.HTTP_201_CREATED,
            False,
            False,
        ),
        (
            "xlsx-template.xlsx",
            models.Template.XLSX_TEMPLATE,
            status.HTTP_201_CREATED,
            True,
            True,
        ),
        (
            "xlsx-template.xlsx",
            models.Template.XLSX_TEMPLATE,
            status.HTTP_401_UNAUTHORIZED,
            True,
            False,
        ),
        (
            "xlsx-template.xlsx",
            models.Template.XLSX_TEMPLATE,
            status.HTTP_201_CREATED,
            True,
            True,
        ),
        (
            "docx-template.docx",
            models.Template.DOCX_TEMPLATE,
            status.HTTP_201_CREATED,
            False,
            False,
        ),
        (
            "docx-template.docx",
            models.Template.DOCX_TEMPLATE,
            status.HTTP_201_CREATED,
            True,
            True,
        ),
        (
            "docx-template.docx",
            models.Template.DOCX_TEMPLATE,
            status.HTTP_401_UNAUTHORIZED,
            True,
            False,
        ),
        (
            "docx-template.docx",
            models.Template.DOCX_TEMPLATE,
            status.HTTP_201_CREATED,
            True,
            True,
        ),
        (
            "docx-mailmerge.docx",
            models.Template.DOCX_MAILMERGE,
            status.HTTP_201_CREATED,
            True,
            True,
        ),
        (
            "docx-mailmerge-syntax.docx",
            models.Template.DOCX_MAILMERGE,
            status.HTTP_400_BAD_REQUEST,
            True,
            True,
        ),
        (
            "docx-template-syntax.docx",
            models.Template.DOCX_TEMPLATE,
            status.HTTP_400_BAD_REQUEST,
            True,
            True,
        ),
        (
            "test.txt",
            models.Template.DOCX_TEMPLATE,
            status.HTTP_400_BAD_REQUEST,
            False,
            False,
        ),
    ],
)
def test_template_create(
    db,
    client,
    admin_client,
    engine,
    template_name,
    status_code,
    require_authentication,
    settings,
    authenticated,
):
    if authenticated:
        client = admin_client

    settings.REQUIRE_AUTHENTICATION = require_authentication

    url = reverse("template-list")

    template_file = django_file(template_name)
    data = {"slug": "test-slug", "template": template_file.file, "engine": engine}
    response = client.post(url, data=data, format="multipart")

    assert response.status_code == status_code

    if status_code == status.HTTP_201_CREATED:
        data = response.json()
        template_link = data["template"]
        response = client.get(template_link)
        assert response.status_code == status.HTTP_200_OK
        file_ = io.BytesIO(response.getvalue())
        if engine == "xlsx-template":
            openpyxl.load_workbook(file_)
        else:
            Document(file_)


@pytest.mark.parametrize(
    "status_code, disable_validation",
    [
        (
            status.HTTP_400_BAD_REQUEST,
            "false",
        ),
        (
            status.HTTP_400_BAD_REQUEST,
            "",
        ),
        (
            status.HTTP_201_CREATED,
            "true",
        ),
    ],
)
def test_disable_validation(
    db,
    status_code,
    admin_client,
    settings,
    disable_validation,
):
    settings.REQUIRE_AUTHENTICATION = False
    url = reverse("template-list")

    template_file = django_file("docx-template-syntax.docx")
    data = {
        "slug": "test-slug",
        "template": template_file.file,
        "engine": models.Template.DOCX_TEMPLATE,
    }
    if disable_validation:
        data["disable_template_validation"] = disable_validation

    response = admin_client.post(url, data=data, format="multipart")
    assert response.status_code == status_code

    if status_code == status.HTTP_201_CREATED:
        data = response.json()
        template_link = data["template"]
        response = admin_client.get(template_link)
        assert response.status_code == status.HTTP_200_OK
        Document(io.BytesIO(response.getvalue()))


@pytest.mark.parametrize(
    "template_name,available_placeholders,sample_data,files,expect_missing_placeholders,engine,status_code",
    [
        (
            "docx-template-placeholdercheck.docx",
            ["foo", "bar", "baz"],
            None,
            [],
            [
                "bar.some_attr",
                "black.png",
                "list",
                "list[]",
                "list[].attribute",
            ],
            models.Template.DOCX_TEMPLATE,
            status.HTTP_400_BAD_REQUEST,
        ),
        (
            "docx-template-placeholdercheck.docx",
            [
                "foo",
                "bar",
                "baz",
                "bar.some_attr",
                "list[].attribute",
                "black.png",
            ],
            None,
            [],
            [],
            models.Template.DOCX_TEMPLATE,
            status.HTTP_201_CREATED,
        ),
        (
            "docx-template-placeholdercheck.docx",
            [
                "foo",
                "bar",
                "baz",
                "bar.some_attr",
                "list[].attribute",
            ],
            None,
            [],
            ["black.png"],
            models.Template.DOCX_TEMPLATE,
            status.HTTP_400_BAD_REQUEST,
        ),
        (
            "docx-template-placeholdercheck.docx",
            None,
            {
                "foo": "hello",
                "bar": {
                    "some_attr": True,
                    "list": [{"attribute": "value"}, {"attribute": "value2"}],
                },
                "baz": "1234",
                "list": [{"attribute": "value"}],
            },
            [django_file("black.png").file],
            [],
            models.Template.DOCX_TEMPLATE,
            status.HTTP_201_CREATED,
        ),
        (
            "docx-template-placeholdercheck.docx",
            None,
            {},
            [django_file("black.png").file],
            [],
            models.Template.DOCX_TEMPLATE,
            status.HTTP_400_BAD_REQUEST,
        ),
        (
            "docx-template-placeholdercheck.docx",
            None,
            {},
            [],
            [],
            models.Template.DOCX_TEMPLATE,
            status.HTTP_201_CREATED,
        ),
        (
            "docx-template-placeholdercheck.docx",
            None,
            {
                "foo": "hello",
                "bar": {
                    "some_attr": True,
                    "list": [{"attribute": "value"}, {"attribute": "value2"}],
                },
                "baz": "1234",
                "list": [{"attribute": "value"}],
            },
            [],
            ["black.png"],
            models.Template.DOCX_TEMPLATE,
            status.HTTP_400_BAD_REQUEST,
        ),
        (
            "docx-template-placeholdercheck.docx",
            None,
            {
                "foo": "hello",
                "bar": {
                    "some_attr": True,
                    "list": [{"attribute": "value"}, {"attribute": "value2"}],
                },
            },
            [django_file("black.png").file],
            ["baz", "list", "list[]", "list[].attribute"],
            models.Template.DOCX_TEMPLATE,
            status.HTTP_400_BAD_REQUEST,
        ),
        (
            "docx-mailmerge.docx",
            None,
            {
                "foo": "hello",
                "bar": {
                    "some_attr": True,
                    "list": [{"attribute": "value"}, {"attribute": "value2"}],
                },
            },
            [],
            ["test"],
            models.Template.DOCX_MAILMERGE,
            status.HTTP_400_BAD_REQUEST,
        ),
        (
            "docx-mailmerge.docx",
            None,
            {
                "foo": "hello",
                "bar": {
                    "some_attr": True,
                    "list": [{"attribute": "value"}, {"attribute": "value2"}],
                },
            },
            [],
            ["test"],
            models.Template.DOCX_MAILMERGE,
            status.HTTP_400_BAD_REQUEST,
        ),
        (
            "docx-mailmerge.docx",
            None,
            {"test": "hello"},
            [],
            [],
            models.Template.DOCX_MAILMERGE,
            status.HTTP_201_CREATED,
        ),
        (
            "docx-mailmerge.docx",
            ["test", "blah"],
            {"test": "hello"},
            [],
            [],
            models.Template.DOCX_MAILMERGE,
            status.HTTP_400_BAD_REQUEST,
        ),
        (
            "docx-mailmerge.docx",
            [],
            {"test": "hello"},
            [django_file("black.png").file],
            [],
            models.Template.DOCX_MAILMERGE,
            status.HTTP_400_BAD_REQUEST,
        ),
        (
            "docx-template-image-placeholder-header-footer.docx",
            ["black.png", "white.png"],
            None,
            [],
            [],
            models.Template.DOCX_TEMPLATE,
            status.HTTP_201_CREATED,
        ),
    ],
)
def test_template_create_with_available_placeholders(
    db,
    admin_client,
    engine,
    template_name,
    available_placeholders,
    sample_data,
    files,
    status_code,
    settings,
    expect_missing_placeholders,
):
    settings.DOCXTEMPLATE_JINJA_EXTENSIONS = ["jinja2.ext.loopcontrols"]
    url = reverse("template-list")

    template_file = django_file(template_name)
    data = {
        "slug": "test-slug",
        "template": template_file.file,
        "files": files,
        "engine": engine,
    }
    if sample_data:
        data["sample_data"] = json.dumps(sample_data)
    if available_placeholders:
        data["available_placeholders"] = available_placeholders

    response = admin_client.post(url, data=data, format="multipart")
    assert response.status_code == status_code, response.json()

    if status_code == status.HTTP_400_BAD_REQUEST:
        resp = response.json()
        expect_missing_str = "; ".join(expect_missing_placeholders)

        if sample_data and available_placeholders:
            # validation only allows one of these two params
            assert (
                resp["non_field_errors"][0]
                == "Only one of available_placeholders and sample_data is allowed"
            )
        elif engine == models.Template.DOCX_MAILMERGE and files:
            assert (
                resp["non_field_errors"][0]
                == 'Files are only accepted with the "docx-template" engine'
            )
        elif not sample_data and files:
            assert (
                resp["non_field_errors"][0]
                == "Files are only accepted when also providing sample_data"
            )
        else:
            # we expect some missing placeholders
            assert resp["non_field_errors"][0] in [
                f"Template uses unavailable placeholders: {expect_missing_str}",
                f'No file for image "{expect_missing_str}" provided!',
            ]

    if status_code == status.HTTP_201_CREATED:
        data = response.json()
        template_link = data["template"]
        response = admin_client.get(template_link)
        assert response.status_code == status.HTTP_200_OK
        Document(io.BytesIO(response.getvalue()))


@pytest.mark.parametrize(
    "template__engine,template__template",
    [(models.Template.DOCX_TEMPLATE, django_file("docx-template-filters.docx"))],
)
@pytest.mark.parametrize(
    "template_name,status_code",
    [
        ("docx-template.docx", status.HTTP_200_OK),
        ("test.txt", status.HTTP_400_BAD_REQUEST),
    ],
)
def test_template_update(db, client, template, template_name, status_code):
    url = reverse("template-detail", args=[template.pk])

    template_file = django_file(template_name)
    data = {"description": "Test description", "template": template_file.file}
    response = client.patch(url, data=data, format="multipart")
    assert response.status_code == status_code

    if status_code == status.HTTP_200_OK:
        assert os.path.isfile(template.template.path) is False
        template.refresh_from_db()
        assert template.description == "Test description"


@pytest.mark.parametrize("template__template", [django_file("docx-template.docx")])
def test_template_destroy(db, client, template):
    url = reverse("template-detail", args=[template.pk])

    response = client.delete(url)
    assert response.status_code == status.HTTP_204_NO_CONTENT
    assert os.path.isfile(template.template.path) is False


@pytest.mark.parametrize(
    "template__slug,template__engine,template__template",
    [
        (
            "TestNameTemplate",
            models.Template.DOCX_TEMPLATE,
            django_file("docx-template.docx"),
        ),
        (
            "TestNameMailMerge",
            models.Template.DOCX_MAILMERGE,
            django_file("docx-mailmerge.docx"),
        ),
    ],
)
def test_template_merge_docx(
    db, client, mock_filefield_name_validation, template, snapshot
):
    url = reverse("template-merge", args=[template.pk])

    response = client.post(url, data={"data": {"test": "Test input"}}, format="json")
    assert response.status_code == status.HTTP_200_OK
    assert (
        response.get("content-type")
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    docx = Document(io.BytesIO(response.getvalue()))
    xml = etree.tostring(docx._element.body, encoding="unicode", pretty_print=True)
    try:
        snapshot.assert_match(xml)
    except AssertionError:  # pragma: no cover
        with open(f"/tmp/{template.slug}.docx", "wb") as output:
            output.write(response.getvalue())
        print("Template output changed. Check file at %s" % output.name)
        raise


@pytest.mark.parametrize(
    "placeholder,template_content",
    [
        ("{{blah}}", {"blah": "blub"}),
        (
            '{{NAME and ", represents " + NAME}}',
            {"NAME": "foo"},
        ),
        (
            '{{NAME and ", represents " + NAME}}',
            {"NAME": ""},
        ),
        # passed data should be escaped
        ("{{escapeme}}", {"escapeme": "<&>"}),
    ],
)
def test_merge_expression(
    docx_template_with_placeholder, client, snapshot, placeholder, template_content
):
    """Test evaluation of some custom expressions.

    Use this test to try out expressions without creating a new docx template for each
    variant.
    """
    template = docx_template_with_placeholder(placeholder)

    url = reverse("template-merge", args=[template.pk])

    response = client.post(url, data={"data": template_content}, format="json")
    assert response.status_code == status.HTTP_200_OK

    docx = Document(io.BytesIO(response.getvalue()))
    xml = etree.tostring(docx._element.body, encoding="unicode", pretty_print=True)
    try:
        snapshot.assert_match(xml)
    except AssertionError:  # pragma: no cover
        with open(f"/tmp/{template.slug}.docx", "wb") as output:
            output.write(response.getvalue())
        print("Template output changed. Check file at %s" % output.name)
        raise


@pytest.mark.parametrize(
    "placeholder,template_content",
    [
        ("{{blah}}", {"blah": "blub"}),
        (
            '{{NAME and ", represents " + NAME}}',
            {"NAME": "foo"},
        ),
        (
            '{{NAME and ", represents " + NAME}}',
            {"NAME": ""},
        ),
    ],
)
@pytest.mark.parametrize(
    "template__engine",
    [models.Template.DOCX_TEMPLATE],
)
def test_validate_expression(
    docx_template_with_placeholder, client, placeholder, template_content
):
    """Test validation of templates with custom expressions."""
    template = docx_template_with_placeholder(placeholder)

    serializer = serializers.TemplateSerializer()
    serializer.instance = template

    serializer.validate({"data": template_content})


# This needs a strange parametrization. If `unoconv_local` is in a separate
# `parametrize()`, the template filename in the second test will be appended with a
# hash and the test fails
@pytest.mark.parametrize(
    "template__engine,template__template,unoconv_local",
    [
        (models.Template.DOCX_TEMPLATE, django_file("docx-template.docx"), True),
        (models.Template.DOCX_TEMPLATE, django_file("docx-template.docx"), False),
    ],
)
def test_template_merge_as_pdf(
    db, settings, unoconv_local, client, mock_filefield_name_validation, template
):
    url = reverse("template-merge", args=[template.pk])

    response = client.post(
        url, data={"data": {"test": "Test input"}, "convert": "pdf"}, format="json"
    )
    assert response.status_code == status.HTTP_200_OK
    assert response["Content-Type"] == "application/pdf"
    assert f"{template.pk}.pdf" in response["Content-Disposition"]
    assert response.content[0:4] == b"%PDF"


@pytest.mark.parametrize(
    "template__engine,template__template",
    [(models.Template.DOCX_TEMPLATE, django_file("docx-template-loopcontrols.docx"))],
)
def test_template_merge_jinja_extensions_docx(
    db, client, mock_filefield_name_validation, template, settings, snapshot
):
    settings.DOCXTEMPLATE_JINJA_EXTENSIONS = ["jinja2.ext.loopcontrols"]

    url = reverse("template-merge", args=[template.pk])

    response = client.post(url, data={"data": {"test": "Test input"}}, format="json")
    assert response.status_code == status.HTTP_200_OK
    assert (
        response.get("content-type")
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    docx = Document(io.BytesIO(response.getvalue()))
    xml = etree.tostring(docx._element.body, encoding="unicode", pretty_print=True)
    snapshot.assert_match(xml)


@pytest.mark.parametrize(
    "missing_file,wrong_mime,status_code",
    [
        (False, False, status.HTTP_200_OK),
        (False, True, status.HTTP_400_BAD_REQUEST),
        (True, False, status.HTTP_400_BAD_REQUEST),
    ],
)
@pytest.mark.parametrize(
    "template__engine,template__template",
    [(models.Template.DOCX_TEMPLATE, django_file("docx-template-filters.docx"))],
)
def test_template_merge_jinja_filters_docx(
    db,
    client,
    mock_filefield_name_validation,
    template,
    snapshot,
    settings,
    tmp_path,
    missing_file,
    wrong_mime,
    status_code,
):
    settings.LANGUAGE_CODE = "de-ch"
    url = reverse("template-merge", args=[template.pk])

    # Couldn't put this into `parametrize`. For some reason, in the second run, the
    # template name is extended with a seemingly random string.
    template.template = django_file("docx-template-filters.docx")
    template.save()

    data = {
        "data": json.dumps(
            {
                "test_date": "1984-09-15",
                "test_time": "23:24",
                "test_datetime": "1984-09-15 23:23",
                "test_datetime2": "23:23-1984-09-15",
                "test_none": None,
                "test_nested": {"multiline": "This is\na test."},
            }
        ),
    }

    if not missing_file:
        file = django_file("black.png").file
        if wrong_mime:
            # create a file with the correct filename (black.png) but with
            # the contents of a docx.
            file = tmp_path / "black.png"
            for line in template.template.file:
                file.write_bytes(line)
            file = file.open("rb")

        data["files"] = [file]

    response = client.post(url, data=data, format="multipart")
    assert response.status_code == status_code

    if status_code == status.HTTP_200_OK:
        assert (
            response.get("content-type")
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        docx = Document(io.BytesIO(response.getvalue()))
        xml = etree.tostring(docx._element.body, encoding="unicode", pretty_print=True)
        snapshot.assert_match(xml)


@pytest.mark.parametrize(
    "template__engine,template__template",
    [(models.Template.DOCX_TEMPLATE, django_file("docx-template-filters.docx"))],
)
@pytest.mark.parametrize(
    "file_value",
    [None, ""],
)
def test_template_merge_file_reset(
    db,
    client,
    mock_filefield_name_validation,
    template,
    settings,
    file_value,
):
    settings.LANGUAGE_CODE = "de-ch"
    url = reverse("template-merge", args=[template.pk])

    # Couldn't put this into `parametrize`. For some reason, in the second run, the
    # template name is extended with a seemingly random string.
    template.template = django_file("docx-template-filters.docx")
    template.save()

    data = {
        "data": {
            "test_date": "1984-09-15",
            "test_time": "23:24",
            "test_datetime": "1984-09-15 23:23",
            "test_datetime2": "23:23-1984-09-15",
            "test_none": None,
            "test_nested": {"multiline": "This is\na test."},
            "black.png": file_value,
        }
    }

    response = client.post(url, data=data, format="json")
    assert response.status_code == status.HTTP_200_OK

    assert (
        response.get("content-type")
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@pytest.mark.parametrize(
    "sample,expected",
    [
        ({"foo": {"bar": ["foo", "blah"]}}, ["foo", "foo.bar", "foo.bar[]"]),
        (
            {
                "this": {
                    "is": {
                        "a": [
                            {
                                "list": {
                                    "with": {
                                        "a": ["nested", "object", "and", "a", "list"]
                                    }
                                }
                            }
                        ]
                    }
                }
            },
            [
                "this",
                "this.is",
                "this.is.a",
                "this.is.a[]",
                "this.is.a[].list",
                "this.is.a[].list.with",
                "this.is.a[].list.with.a",
                "this.is.a[].list.with.a[]",
            ],
        ),
    ],
)
def test_sample_to_placeholders(sample, expected):
    ts = serializers.TemplateSerializer()
    assert ts._sample_to_placeholders(sample) == sorted(expected)


@pytest.mark.parametrize(
    "template__engine,template__template",
    [
        (
            models.Template.DOCX_TEMPLATE,
            django_file("docx-template-placeholdercheck.docx"),
        )
    ],
)
def test_template_merge_missing_data(
    db, client, mock_filefield_name_validation, template, settings
):
    settings.DOCXTEMPLATE_JINJA_EXTENSIONS = ["jinja2.ext.loopcontrols"]

    url = reverse("template-merge", args=[template.pk])

    response = client.post(url, data={"data": {"blah": "Test input"}}, format="json")

    assert response.status_code == status.HTTP_400_BAD_REQUEST
    assert response.json() == [
        "Placeholder from template not found in data: 'bar' is undefined"
    ]
